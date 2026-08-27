"""抽樣 208會議/20830 高風險檔，只讀欄位名稱、不取任何儲存格值。

輸出：每個檔命中哪些個資欄位類型 + 全樣本彙總。
下載的暫存檔跑完即刪。
"""
import csv, io, os, random, re, subprocess, sys, shutil
from collections import Counter
from pathlib import Path

PROJ = Path("/Users/liaoneil/projects/emei/drive-inventory")
sys.path.insert(0, str(PROJ / "scripts"))
os.chdir(PROJ)
from auth import get_credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload

TMP = Path("/private/tmp/claude-501/-Users-liaoneil-projects-Neil-agent/1af11a15-067c-4764-aa4c-3912539adf49/scratchpad/pii-sample")

# 個資欄位詞表：出現在「欄名」才算
FIELDS = {
    "姓名": ["姓名", "名字", "學員姓名", "班員姓名"],
    "身分證": ["身分證", "身份證", "統一編號", "證號"],
    "電話/手機": ["電話", "手機", "聯絡方式", "行動"],
    "地址": ["地址", "住址", "戶籍"],
    "生日/年齡": ["生日", "出生", "年齡", "歲"],
    "性別": ["性別", "乾坤", "乾道", "坤道"],
    "Email": ["email", "e-mail", "信箱", "電子郵件"],
    "緊急聯絡人": ["緊急", "監護", "家長"],
    "健康/飲食": ["血型", "病史", "過敏", "素食", "葷素", "疾病", "健康"],
    "學校/職業": ["學校", "就讀", "年級", "職業", "服務單位"],
    "道場身分": ["佛堂", "道號", "求道", "點傳", "引保", "壇主", "單位", "區", "組"],
}

def classify(labels):
    hit = set()
    for lab in labels:
        low = lab.lower()
        for field, kws in FIELDS.items():
            if any(k.lower() in low for k in kws):
                hit.add(field)
    return hit

def labels_xlsx(p):
    import openpyxl
    out = []
    wb = openpyxl.load_workbook(p, read_only=True, data_only=True)
    for ws in wb.worksheets[:3]:
        for i, row in enumerate(ws.iter_rows(max_row=12, max_col=40, values_only=True)):
            for c in row:
                if isinstance(c, str) and 1 <= len(c.strip()) <= 20:
                    out.append(c.strip())
            if i > 12: break
    wb.close()
    return out

def labels_xls(p):
    import xlrd
    out = []
    bk = xlrd.open_workbook(p)
    for sh in bk.sheets()[:3]:
        for r in range(min(12, sh.nrows)):
            for c in range(min(40, sh.ncols)):
                v = sh.cell_value(r, c)
                if isinstance(v, str) and 1 <= len(v.strip()) <= 20:
                    out.append(v.strip())
    return out

def labels_docx(p):
    import docx
    d = docx.Document(p)
    out = []
    for t in d.tables[:5]:
        for row in t.rows[:12]:
            for cell in row.cells:
                s = cell.text.strip()
                if 1 <= len(s) <= 20:
                    out.append(s)
    for para in d.paragraphs[:40]:
        s = para.text.strip()
        if 1 <= len(s) <= 30:
            out.append(s)
    return out

def labels_doc(p):
    r = subprocess.run(["textutil", "-convert", "txt", "-stdout", str(p)],
                       capture_output=True, timeout=60)
    txt = r.stdout.decode("utf-8", "ignore")
    return [s.strip() for s in re.split(r"[\n\t　 ]+", txt) if 1 <= len(s.strip()) <= 20][:600]

rows = []
with open("data/pii-scan/candidates.csv", encoding="utf-8-sig") as f:
    for r in csv.DictReader(f):
        if r["資料夾路徑"].startswith("208會議/20830") and r["風險級"].startswith("高"):
            rows.append(r)
random.seed(20260827)
sample = random.sample(rows, 20)

svc = build("drive", "v3", credentials=get_credentials())
agg = Counter()
per_file = []
for r in sample:
    fid, name = r["檔案id"], r["檔名"]
    ext = name.rsplit(".", 1)[-1].lower()
    dest = TMP / f"{fid}.{ext}"
    try:
        req = svc.files().get_media(fileId=fid, supportsAllDrives=True)
        buf = io.BytesIO()
        dl = MediaIoBaseDownload(buf, req)
        done = False
        while not done:
            _, done = dl.next_chunk()
        dest.write_bytes(buf.getvalue())
        labs = {"xlsx": labels_xlsx, "xls": labels_xls,
                "docx": labels_docx, "doc": labels_doc}[ext](dest)
        hits = classify(labs)
    except Exception as e:
        per_file.append((name, f"讀取失敗: {type(e).__name__}"))
        continue
    finally:
        if dest.exists(): dest.unlink()
    agg.update(hits)
    per_file.append((name, "、".join(sorted(hits)) or "（無個資欄位）"))

print(f"抽樣 {len(sample)} 檔（208會議/20830 高風險 496 筆隨機抽）\n")
print("=== 逐檔命中欄位 ===")
for n, h in per_file:
    print(f"  {n[:46]:<48} {h}")
print("\n=== 彙總：20 檔中有幾檔含此欄位 ===")
for field, cnt in agg.most_common():
    print(f"  {field:<12} {cnt}/20")
shutil.rmtree(TMP, ignore_errors=True)
